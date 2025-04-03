import streamlit as st
import requests
import os
import json
import base64
from pathlib import Path
from mistralai import Mistral, DocumentURLChunk, ImageURLChunk, TextChunk
from mistralai.models import OCRResponse
from PIL import Image
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime

# Set page configuration
st.set_page_config(
    page_title="Mistral OCR Document Processor",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.title("Mistral OCR Document Processor")
st.write("Upload a PDF or image file (JPEG, PNG) or take a photo to process with Mistral's OCR service")


# Function to replace image placeholders in markdown with base64-encoded images
def replace_images_in_markdown(markdown_str: str, images_dict: dict) -> str:
    """
    Replace image placeholders in markdown with base64-encoded images.

    Args:
        markdown_str: Markdown text containing image placeholders
        images_dict: Dictionary mapping image IDs to base64 strings

    Returns:
        Markdown text with images replaced by base64 data
    """
    for img_name, base64_str in images_dict.items():
        markdown_str = markdown_str.replace(
            f"![{img_name}]({img_name})", f"![{img_name}]({base64_str})"
        )
    return markdown_str


# Function to get_combined_markdown with pagination support
def get_combined_markdown(ocr_response: OCRResponse, page_start: int = 0, page_end: int = None) -> tuple[str, dict]:
    """
    Combine OCR text and images into a single markdown document with pagination support.
    
    Args:
        ocr_response: Response from OCR processing containing text and images
        page_start: Starting page index (0-based)
        page_end: Ending page index (exclusive), None for all pages
        
    Returns:
        Tuple of (markdown string, image data dictionary)
    """
    markdowns: list[str] = []
    image_data = {}
    
    # Calculate page range
    if page_end is None:
        page_end = len(ocr_response.pages)
    pages = ocr_response.pages[page_start:page_end]
    
    # Extract images from page
    for page in pages:
        for img in page.images:
            image_data[img.id] = img.image_base64
        # Replace image placeholders with actual images
        markdowns.append(replace_images_in_markdown(page.markdown, image_data))

    return "\n\n".join(markdowns), image_data


# Function to process document with pagination
def process_document_with_pagination(client, document_response, page_size=5):
    """
    Process document with pagination to handle large documents efficiently.
    
    Args:
        client: Mistral client
        document_response: Initial document response
        page_size: Number of pages to process at once
        
    Returns:
        Generator yielding (markdown, image_data) tuples
    """
    total_pages = len(document_response.pages)
    for i in range(0, total_pages, page_size):
        markdown, image_data = get_combined_markdown(
            document_response,
            page_start=i,
            page_end=min(i + page_size, total_pages)
        )
        yield markdown, image_data


# Function to display paginated content
def display_paginated_content(markdown_text, image_data, page_num, total_pages):
    """
    Display paginated content with navigation controls.
    
    Args:
        markdown_text: Markdown text to display
        image_data: Dictionary of image data
        page_num: Current page number
        total_pages: Total number of pages
    """
    # Create navigation controls
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if page_num > 1:
            if st.button("Previous Page"):
                st.session_state.current_page -= 1
                st.rerun()
    with col2:
        st.write(f"Page {page_num} of {total_pages}")
    with col3:
        if page_num < total_pages:
            if st.button("Next Page"):
                st.session_state.current_page += 1
                st.rerun()
    
    # Display current page content
    st.markdown(markdown_text, unsafe_allow_html=True)


# Function to export to Word with pagination
def export_to_word_paginated(markdown_texts: list[str]) -> bytes:
    """
    Export paginated markdown text to a Word document.
    
    Args:
        markdown_texts: List of markdown texts for each page
        
    Returns:
        Bytes of the Word document
    """
    doc = Document()
    
    for markdown_text in markdown_texts:
        # Remove base64 image references from text while keeping other content
        clean_text = re.sub(r'!\[.*?\]\(data:image/[^;]+;base64,[^\)]+\)', '', markdown_text)
        
        # Clean math expressions
        clean_text = re.sub(r'\$\\mathbf\{([^}]+)\}\$', r'\1', clean_text)
        
        # Extract content parts (text and tables)
        content_parts = extract_tables_from_markdown(clean_text)
        
        for part in content_parts:
            if part["type"] == "text":
                # Add text paragraphs
                paragraphs = part["content"].split('\n')
                for p in paragraphs:
                    if p.strip():
                        # Check for headers
                        header_match = re.match(r'^(#{1,6})\s+(.+)$', p.strip())
                        if header_match:
                            level = len(header_match.group(1))
                            text = header_match.group(2)
                            doc.add_heading(text, level=level)
                        else:
                            # Process bold text
                            parts = re.split(r'(\*\*.*?\*\*)', p.strip())
                            if len(parts) > 1:
                                paragraph = doc.add_paragraph()
                                for part_text in parts:
                                    if part_text.startswith('**') and part_text.endswith('**'):
                                        run = paragraph.add_run(part_text[2:-2])
                                        run.bold = True
                                    else:
                                        paragraph.add_run(part_text)
                            else:
                                doc.add_paragraph(p.strip())
            
            elif part["type"] == "table":
                if part["headers"] and part["data"]:
                    table = doc.add_table(rows=1, cols=len(part["headers"]))
                    table.style = 'Table Grid'
                    
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(part["headers"]):
                        run = header_cells[i].paragraphs[0].add_run(header)
                        run.bold = True
                    
                    for row_data in part["data"]:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row_data):
                            cell_text = cell.strip()
                            if cell_text.startswith('$'):
                                paragraph = row_cells[i].paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                if 'CR' in cell_text:
                                    cell_text = f"({cell_text.replace('CR', '').strip()})"
                            row_cells[i].text = cell_text
                            if i == 0 and cell_text.replace('.', '').isdigit():
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save the document to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


# Function to convert image to base64
def image_to_base64(image):
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')


# Function to extract tables from markdown
def extract_tables_from_markdown(markdown_text):
    """Extract tables and surrounding text from markdown."""
    if not markdown_text:
        return []
        
    # Split the markdown by table markers
    parts = re.split(r'(\n\|.*\|.*\n(?:\|.*\|.*\n)+)', markdown_text)

    result = []
    i = 0
    while i < len(parts):
        if i < len(parts) and not parts[i].strip().startswith('|'):
            # This is text content
            if parts[i].strip():
                result.append({"type": "text", "content": parts[i].strip()})

        # Check if we have a table part
        if i < len(parts) and '|' in parts[i]:
            try:
                table_text = parts[i].strip()
                # Extract table rows
                rows = [row.strip() for row in table_text.split('\n') if row.strip().startswith('|')]

                if not rows:
                    # If no valid rows found, treat as text
                    if table_text.strip():
                        result.append({"type": "text", "content": table_text.strip()})
                    i += 1
                    continue

                # Skip separator row (contains :--:, :-- or --:)
                header_rows = []
                data_rows = []
                separator_found = False

                for j, row in enumerate(rows):
                    if ':--' in row or '--:' in row or '---' in row:
                        header_rows = rows[:j]
                        data_rows = rows[j + 1:]
                        separator_found = True
                        break

                if not separator_found:
                    # No separator found, treat first row as header
                    if len(rows) > 0:
                        header_rows = [rows[0]]
                        data_rows = rows[1:] if len(rows) > 1 else []
                    else:
                        # If no rows at all, skip this part
                        i += 1
                        continue

                # Parse headers safely
                headers = []
                if header_rows:
                    header_cells = header_rows[0].split('|')
                    # Remove empty cells from start and end
                    header_cells = [cell for cell in header_cells if cell.strip()]
                    headers = [cell.strip() for cell in header_cells]

                # Parse data rows safely
                data = []
                for row in data_rows:
                    cells = row.split('|')
                    # Remove empty cells from start and end
                    cells = [cell for cell in cells if cell.strip()]
                    if cells:  # Only add non-empty rows
                        # Ensure each row has the same number of columns as headers
                        row_data = [cell.strip() for cell in cells]
                        if headers:
                            # Pad or truncate row to match header length
                            if len(row_data) < len(headers):
                                row_data.extend([''] * (len(headers) - len(row_data)))
                            elif len(row_data) > len(headers):
                                row_data = row_data[:len(headers)]
                        data.append(row_data)

                if headers or data:  # Only add if we have either headers or data
                    result.append({
                        "type": "table",
                        "headers": headers,
                        "data": data
                    })
            except Exception as e:
                # If table parsing fails, treat as text
                if parts[i].strip():
                    result.append({"type": "text", "content": parts[i].strip()})

        i += 1

    return result


# Get API key from secrets or user input
def get_api_key():
    # Try to get API key from secrets
    try:
        return st.secrets["mistral"]["api_key"]
    except:
        # If not available in secrets, return empty string
        return ""


# API key input with default from secrets if available
default_api_key = get_api_key()
user_api_key = st.text_input(
    "Enter your Mistral API key:",
    value="",
    type="password",
    help="Enter your Mistral API key or configure it in st.secrets['mistral_api_key']"
)

# Use the provided API key or fall back to secrets
api_key = user_api_key if user_api_key else default_api_key

if not api_key:
    st.warning("Please provide a Mistral API key to use this application.")
    st.stop()

# Create tabs for different input methods
input_tab1, input_tab2 = st.tabs(["Upload Document", "Take Photo"])

with input_tab1:
    # File uploader with size limit of 50MB
    uploaded_file = st.file_uploader("Choose a PDF or image file (max 50MB)", type=["pdf", "jpeg", "jpg", "png"])

    if uploaded_file is not None:
        # Check file size (50MB = 50 * 1024 * 1024 bytes)
        file_size_limit = 50 * 1024 * 1024  # 50MB in bytes
        file_size = len(uploaded_file.getvalue())

        if file_size > file_size_limit:
            st.error(f"File size exceeds the 50MB limit. Your file is {file_size / (1024 * 1024):.2f}MB.")
        else:
            st.success(f"File uploaded: {uploaded_file.name} ({file_size / (1024 * 1024):.2f}MB)")

            # Process button
            if st.button("Process Document with OCR", key="process_document"):
                try:
                    with st.spinner(f"Processing {uploaded_file.name.split('.')[-1].upper()} with Mistral OCR..."):
                        # Initialize Mistral client
                        client = Mistral(api_key=api_key)

                        # Process document with OCR based on file type
                        file_extension = uploaded_file.name.split('.')[-1].lower()
                        if file_extension in ['jpeg', 'jpg', 'png']:
                            # For images, convert to base64 directly from memory
                            base64_image = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')

                            # Process image with OCR
                            document_response = client.ocr.process(
                                document=ImageURLChunk(image_url=f"data:image/{file_extension};base64,{base64_image}"),
                                model="mistral-ocr-latest",
                                include_image_base64=True
                            )
                        else:
                            # For PDFs, upload directly from memory
                            mistral_uploaded_file = client.files.upload(
                                file={
                                    "file_name": uploaded_file.name,
                                    "content": uploaded_file.getvalue(),
                                },
                                purpose="ocr",
                            )

                            # Get URL for the uploaded file
                            signed_url = client.files.get_signed_url(file_id=mistral_uploaded_file.id, expiry=1)

                            # Process document with OCR, including embedded images
                            document_response = client.ocr.process(
                                document=DocumentURLChunk(document_url=signed_url.url),
                                model="mistral-ocr-latest",
                                include_image_base64=True
                            )

                        # Initialize session state for pagination
                        if 'current_page' not in st.session_state:
                            st.session_state.current_page = 1
                        if 'total_pages' not in st.session_state:
                            st.session_state.total_pages = len(document_response.pages)
                        if 'document_response' not in st.session_state:
                            st.session_state.document_response = document_response
                        if 'all_markdowns' not in st.session_state:
                            st.session_state.all_markdowns = []

                        # Process current page
                        page_start = (st.session_state.current_page - 1) * 5
                        page_end = min(page_start + 5, st.session_state.total_pages)
                        markdown, image_data = get_combined_markdown(
                            document_response,
                            page_start=page_start,
                            page_end=page_end
                        )

                        # Store results in session state
                        st.session_state.ocr_results = markdown
                        st.session_state.show_results = True

                        st.success("Document processing completed!")

                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    st.session_state.show_results = False

            # Display OCR results if available
            if hasattr(st.session_state, 'show_results') and st.session_state.show_results:
                # Display results
                st.subheader("OCR Results")

                # Create tabs for original text and markdown
                tab1, tab2 = st.tabs(["Original Text", "Markdown"])

                with tab1:
                    # Add export button in a columns layout to save space
                    col1, col2 = st.columns([1, 4])
                    with col1:
                        if st.button("Export to Word", key="export_doc"):
                            with st.spinner("Generating Word document..."):
                                try:
                                    # Process all pages for export
                                    all_markdowns = []
                                    for i in range(0, st.session_state.total_pages, 5):
                                        markdown, _ = get_combined_markdown(
                                            st.session_state.document_response,
                                            page_start=i,
                                            page_end=min(i + 5, st.session_state.total_pages)
                                        )
                                        all_markdowns.append(markdown)
                                    
                                    output_filename = f"ocr_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                    doc_bytes = export_to_word_paginated(all_markdowns)
                                    st.download_button(
                                        label="📥 Download Document",
                                        data=doc_bytes,
                                        file_name=output_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    )
                                except Exception as e:
                                    st.error(f"Error exporting to Word: {str(e)}")
                    
                    # Display paginated content
                    display_paginated_content(
                        st.session_state.ocr_results,
                        {},
                        st.session_state.current_page,
                        st.session_state.total_pages
                    )

                with tab2:
                    # Display raw markdown with syntax highlighting
                    st.code(st.session_state.ocr_results, language="markdown")

with input_tab2:
    st.write("Take a photo with your camera")

    # Camera input
    camera_image = st.camera_input("Take a picture")

    if camera_image is not None:
        # Process button for camera image
        if st.button("Process Image with OCR", key="process_image"):
            try:
                with st.spinner("Processing image with Mistral OCR..."):
                    # Initialize Mistral client
                    client = Mistral(api_key=api_key)

                    # Convert the camera image to base64
                    base64_image = base64.b64encode(camera_image.getvalue()).decode('utf-8')

                    # Process image with OCR using ImageURLChunk with image_url
                    image_response = client.ocr.process(
                        document=ImageURLChunk(image_url=f"data:image/jpeg;base64,{base64_image}"),
                        model="mistral-ocr-latest",
                        include_image_base64=True
                    )

                    # Get combined markdown
                    combined_markdown = get_combined_markdown(image_response)

                    # Store results in session state
                    st.session_state.ocr_results = combined_markdown
                    st.session_state.show_results = True

                    st.success("Image processing completed!")

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.session_state.show_results = False

        # Display OCR results if available
        if hasattr(st.session_state, 'show_results') and st.session_state.show_results:
            # Display results
            st.subheader("OCR Results")

            # Create tabs for original text and markdown
            tab1, tab2 = st.tabs(["Original Text", "Markdown"])

            with tab1:
                # Add export button in a columns layout to save space
                col1, col2 = st.columns([1, 4])
                with col1:
                    if st.button("Export to Word", key="export_img"):
                        with st.spinner("Generating Word document..."):
                            try:
                                output_filename = f"ocr_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                doc_bytes = export_to_word(st.session_state.ocr_results)
                                st.download_button(
                                    label="📥 Download Document",
                                    data=doc_bytes,
                                    file_name=output_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                            except Exception as e:
                                st.error(f"Error exporting to Word: {str(e)}")
                
                # Display combined markdowns and images
                st.markdown(st.session_state.ocr_results, unsafe_allow_html=True)

            with tab2:
                # Display raw markdown with syntax highlighting
                st.code(st.session_state.ocr_results, language="markdown")

# Add some information about the app
st.sidebar.title("About")
st.sidebar.info(
    "This app uses Mistral AI's OCR service to extract text and images from PDF documents, "
    "image files (JPEG, PNG), or photos taken with your camera. Upload a document or take a photo, "
    "click 'Process', and view the extracted content in markdown format."
)

# Add API key configuration information
st.sidebar.title("API Key Configuration")
st.sidebar.info(
    "You can provide your Mistral API key in two ways:\n"
    "1. Enter it directly in the text field\n"
    "2. Set it in your Streamlit secrets.toml file as:\n"
    "```\n"
    "mistral_api_key = 'your-api-key-here'\n"
    "```"
)

# Add requirements information
st.sidebar.title("Requirements")
st.sidebar.code("pip install mistralai pillow")